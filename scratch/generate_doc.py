import urllib.request
import base64
import os
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_documentation():
    doc = docx.Document()
    
    # Page setup (Standard 1-inch margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Fonts
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Helper colors
    COLOR_PRIMARY = RGBColor(0x00, 0x70, 0xC0)  # Deep Ocean Blue
    COLOR_SECONDARY = RGBColor(0x00, 0x99, 0x99) # Teal
    COLOR_DARK = RGBColor(0x1F, 0x29, 0x37)     # Dark Slate
    COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)    # Gray

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = COLOR_MUTED
        p.paragraph_format.space_after = Pt(18)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_DARK
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_callout(text, label=""):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        # Border & background styling
        shading = parse_xml(r'<w:shd {} w:fill="F0F9FF"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        if label:
            r_lbl = p.add_run(label + "\n")
            r_lbl.bold = True
            r_lbl.font.color.rgb = COLOR_PRIMARY
        r_txt = p.add_run(text)
        r_txt.font.italic = True
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- DOCUMENT HEADER / PLACEHOLDERS ---
    add_title("VoxLoop System Documentation")
    add_subtitle("In-Flight Multi-Agent Critique & Self-Correction Voice Engine for Support QA")

    # Callout box for User URLs
    callout_p = doc.add_paragraph()
    callout_p.paragraph_format.space_after = Pt(14)
    tbl_url = doc.add_table(rows=1, cols=1)
    tbl_url.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_url = tbl_url.cell(0, 0)
    cell_url.width = Inches(6.5)
    shd_url = parse_xml(r'<w:shd {} w:fill="FFFBEB"/>'.format(nsdecls('w')))
    cell_url._tc.get_or_add_tcPr().append(shd_url)
    
    p_url = cell_url.paragraphs[0]
    p_url.paragraph_format.space_before = Pt(8)
    p_url.paragraph_format.space_after = Pt(8)
    
    r_hdr = p_url.add_run("📌 PROJECT LINKS & REPOSITORY URLS\n")
    r_hdr.bold = True
    r_hdr.font.size = Pt(11)
    r_hdr.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    
    r_gh = p_url.add_run("🔗 GitHub Repository URL: ")
    r_gh.bold = True
    p_url.add_run("[ INSERT YOUR GITHUB REPOSITORY URL HERE ]\n\n")
    
    r_live = p_url.add_run("🌐 Live Deployment URL: ")
    r_live.bold = True
    p_url.add_run("https://voxloop.duckdns.org  (or [ INSERT YOUR LIVE DEPLOYED URL HERE ])")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- 1. EXECUTIVE SUMMARY ---
    add_h1("1. Executive Summary")
    doc.add_paragraph(
        "VoxLoop is a state-of-the-art real-time voice assistant platform engineered around an in-flight multi-agent critique "
        "and self-correction feedback loop. Built using FastAPI, WebSockets, LangGraph, LangChain, and Mistral AI, VoxLoop captures "
        "raw microphone speech from a browser interface, transcribes it locally using faster-whisper, drafts an initial response, "
        "critiques that response against a secondary evaluation LLM across 6 compliance dimensions, rewrites the phrasing using structured "
        "critique feedback, and streams synthesized speech back to the client."
    )

    # --- 2. PROBLEM STATEMENT & TARGET ROLE ---
    add_h1("2. Target Role & Problem Statement")
    add_h2("The Target Role")
    doc.add_paragraph(
        "The primary user role for VoxLoop is Sarah Jenkins, a Quality Assurance (QA) Lead at a FinTech Call Center overseeing 50 support agents."
    )
    add_h2("The Specific Problem")
    doc.add_paragraph(
        "In high-stakes phone customer service (e.g., banking, technical support, insurance), support reps frequently give answers "
        "that are technically correct but fail on critical soft skills: empathy, clarity, or strict regulatory policy compliance. "
        "QA leads currently spend 20+ hours every week listening to past call recordings post-mortem to fill out evaluation forms manually. "
        "Because this evaluation happens hours or days after the call, agents miss out on real-time coaching."
    )
    add_h2("How VoxLoop Solves It")
    doc.add_paragraph(
        "VoxLoop operates as an in-flight live critique & training copilot. It automatically scores support responses across 6 key metrics "
        "(Accuracy, Relevance, Empathy, Clarity, Policy Compliance, and Overall Quality), provides immediate actionable suggestions, "
        "and demonstrates the exact improved phrasing that should have been spoken."
    )

    # --- 3. SYSTEM ARCHITECTURE & DIAGRAM ---
    add_h1("3. System Architecture & Workflow Diagram")
    doc.add_paragraph(
        "The system separates concerns into a clean 4-stage pipeline: Client Interaction, Speech & Storage, LangGraph Multi-Agent Engine, "
        "and LLM Provider."
    )

    # Try downloading or rendering the workflow diagram
    diagram_path = Path("scratch/workflow_diagram.png")
    diagram_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Generate diagram using matplotlib if available or download mermaid ink
    try:
        import matplotlib.pyplot as plt
        import matplotlib.patches as patches

        fig, ax = plt.subplots(figsize=(10, 4.5), dpi=300)
        ax.set_facecolor('#0B0F17')
        fig.patch.set_facecolor('#0B0F17')
        ax.axis('off')

        # Box drawing function
        def draw_box(x, y, w, h, title, subtitle, color, border_color):
            rect = patches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1", 
                                          facecolor=color, edgecolor=border_color, linewidth=1.5)
            ax.add_patch(rect)
            ax.text(x + w/2, y + h*0.65, title, color='white', weight='bold', fontsize=9, ha='center', va='center')
            ax.text(x + w/2, y + h*0.3, subtitle, color='#94A3B8', fontsize=7.5, ha='center', va='center')

        # Arrow function
        def draw_arrow(x1, y1, x2, y2, label=""):
            ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                        arrowprops=dict(arrowstyle="->", color='#00F2FE', lw=1.2, mutation_scale=12))
            if label:
                ax.text((x1+x2)/2, (y1+y2)/2 + 0.15, label, color='#4BC6B9', fontsize=7, ha='center', va='center')

        # 4 Stage Columns
        draw_box(0.2, 2.0, 1.8, 1.2, "1. Client UI", "Next.js 15 & WebSockets", "#141E30", "#00F2FE")
        draw_box(2.5, 2.0, 1.8, 1.2, "2. Speech Engine", "Faster-Whisper & pyttsx3", "#141E30", "#4BC6B9")
        draw_box(4.8, 2.0, 2.2, 1.2, "3. LangGraph Loop", "Draft -> Critique -> Improve", "#141E30", "#FF7E5F")
        draw_box(7.5, 2.0, 1.8, 1.2, "4. LLM Provider", "Mistral AI API", "#141E30", "#9D4EDD")

        # Database box below
        draw_box(2.5, 0.4, 1.8, 0.9, "SQLite Database", "voxloop.db Session Memory", "#141E30", "#3B82F6")

        # Connections
        draw_arrow(2.0, 2.6, 2.5, 2.6, "WebM Audio")
        draw_arrow(4.3, 2.6, 4.8, 2.6, "Transcript")
        draw_arrow(7.0, 2.6, 7.5, 2.6, "JSON Schema")
        draw_arrow(3.4, 2.0, 3.4, 1.3, "Persist Memory")

        plt.title("VoxLoop Multi-Agent Real-Time Architecture", color='white', fontsize=12, pad=12, weight='bold')
        plt.tight_layout()
        plt.savefig(diagram_path, bbox_inches='tight', facecolor=fig.get_facecolor(), edgecolor='none')
        plt.close()
        
        doc.add_picture(str(diagram_path), width=Inches(6.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Figure 1: VoxLoop 4-Stage Multi-Agent Architecture Diagram")
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = COLOR_MUTED
    except Exception as exc:
        print(f"Diagram generation warning: {exc}")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- 4. TECH STACK & SYSTEM COMPONENTS ---
    add_h1("4. Technology Stack & Component Specifications")
    
    tbl_tech = doc.add_table(rows=1, cols=3)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_tech.autofit = False
    
    hdr_cells = tbl_tech.rows[0].cells
    hdr_titles = ["Layer / Module", "Technology Used", "Role in Architecture"]
    widths = [Inches(1.8), Inches(2.0), Inches(2.7)]
    
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].width = widths[i]
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = parse_xml(r'<w:shd {} w:fill="0070C0"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)

    tech_data = [
        ("Frontend Interface", "Next.js 15, React 19, CSS Glassmorphism", "Dual-column layout, live microphone capture, score progress bars, chat bubbles."),
        ("Backend Framework", "FastAPI, Uvicorn, WebSockets", "Async event loop, real-time WebSocket protocol (`/ws/voice-turn`), REST API routes."),
        ("Agent Orchestration", "LangGraph, LangChain", "State-graph execution chaining draft_response -> critique_response -> improve_response."),
        ("LLM Engine", "Mistral AI (mistral-small-latest)", "Primary response generation and structured JSON output evaluation."),
        ("Speech-to-Text (STT)", "faster-whisper", "Local CTranslate2 Whisper model with Voice Activity Detection (VAD) running on CPU/INT8."),
        ("Text-to-Speech (TTS)", "pyttsx3", "Offline cross-platform SAPI / espeak voice synthesis engine saving .wav files."),
        ("Database Layer", "SQLite + SQLAlchemy ORM", "Persistent store for conversation runs, session history, and scorecards (`voxloop.db`)."),
        ("Containerization", "Docker & Docker Compose", "Multi-stage builds packaging backend and frontend into isolated containers."),
        ("Reverse Proxy & SSL", "Nginx + Let's Encrypt (Certbot)", "Production HTTPS/WSS proxying with automated SSL certificate renewal.")
    ]

    for row_idx, (layer, tech, role) in enumerate(tech_data):
        row_cells = tbl_tech.add_row().cells
        for col_idx, text in enumerate([layer, tech, role]):
            row_cells[col_idx].width = widths[col_idx]
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            # Alternating background shading
            if row_idx % 2 == 1:
                shd = parse_xml(r'<w:shd {} w:fill="F9FAFB"/>'.format(nsdecls('w')))
                row_cells[col_idx]._tc.get_or_add_tcPr().append(shd)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- 5. CODEBASE FILES DEEP DIVE ---
    add_h1("5. Python Backend File Structure & Detailed Explanation")
    
    files_info = [
        ("config.py", "Manages application settings, model defaults, and environment variables loaded from backend/.env using Pydantic BaseSettings."),
        ("models.py", "Defines Pydantic data transfer objects (DTOs) for VoiceTurnRequest, GradingScorecard (0-100 bounds), CritiqueResult, and VoiceTurnResponse."),
        ("db.py", "Defines the SQLAlchemy ConversationRun ORM table, engine creation with directory safety, and session lifecycle helper functions."),
        ("audio.py", "Handles asynchronous speech-to-text (transcribe_audio using faster-whisper) and text-to-speech (synthesize_speech using pyttsx3 with Linux exception safety)."),
        ("workflow.py", "Constructs the LangGraph state machine (build_workflow) chaining draft_response, critique_response, and improve_response nodes."),
        ("main.py", "The central FastAPI application providing /health, /api/reset-session, /api/voice-turn, and the real-time WebSocket connection handler (/ws/voice-turn).")
    ]

    for filename, desc in files_info:
        add_h2(f"File: backend/app/{filename}")
        doc.add_paragraph(desc)

    # --- 6. TECHNICAL & ARCHITECTURAL DEEP DIVES ---
    add_h1("6. Technical & Architectural Deep Dives")

    add_h2("01. Why could no one have made this project in 2023?")
    doc.add_paragraph(
        "In early 2023, LLM interactions were restricted to single-pass prompt chains. LangGraph (released in late 2023/2024) "
        "introduced stateful cyclic graphs allowing state to flow through discrete node functions with explicit branch controls. "
        "Additionally, deterministic JSON schema enforcement via Pydantic output parsing and fast edge local STT (faster-whisper) "
        "made real-time feedback loops viable without high API latency or regex output parsing bugs."
    )

    add_h2("02. What is the non-obvious hard part?")
    doc.add_paragraph(
        "Managing Multi-Turn Context Consistency Across a Cyclic Multi-Agent Graph Without Hallucination or Feedback Loops. "
        "If you naive-append the raw Critique object into subsequent turns as conversation history, the LLM starts hallucinating "
        "about its own meta-critique in future turns ('As a critic agent, I scored myself 85...') rather than speaking naturally. "
        "To solve this, VoxLoop enforces strict memory isolation: only the Human Utterance and the Final Improved Response are saved to "
        "long-term SQLite history, while critique scores remain isolated strictly to the active turn evaluation graph."
    )

    add_h2("03. What did you build versus what did the API give you?")
    doc.add_paragraph(
        "The APIs provided raw text completions (Mistral API), audio bytes to text (Faster-Whisper), and string to audio file (pyttsx3). "
        "WE built the entire custom software engine: the real-time WebSocket state machine, non-blocking asyncio thread pools, "
        "the 3-node LangGraph orchestration graph, SQLite multi-turn context formatting (_history_text), session reset handlers, "
        "and the dual-column Next.js glassmorphism frontend."
    )

    add_h2("04. Why does this break if you remove the AI?")
    doc.add_paragraph(
        "VoxLoop is fundamentally an AI-driven evaluation system. Without the Critic Agent LLM, you cannot calculate dynamic scores "
        "for subjective human soft skills like Empathy, Relevance, or Clarity. Static rule matchers can only look for rigid keywords. "
        "Without the Primary Agent LLM, the system cannot take unstructured critique recommendations and automatically synthesize a "
        "rewritten response. Without AI, the app collapses into a simple audio file recorder."
    )

    add_h2("05. What breaks at ten thousand concurrent users?")
    doc.add_paragraph(
        "1. CPU Bottleneck: Local faster-whisper STT and pyttsx3 CPU inference on the FastAPI host will saturate CPU cores. Solution: Offload STT/TTS to a distributed Ray/Celery GPU worker cluster.\n"
        "2. SQLite Write Locks: Simultaneous writes will throw 'database is locked' errors. Solution: Migrate SQLite to PostgreSQL / AWS Aurora with Redis session caching.\n"
        "3. WebSocket File Descriptors: A single FastAPI instance will exhaust file descriptors. Solution: Load balance multiple FastAPI instances behind Nginx with a Redis Pub/Sub backplane."
    )

    # --- 7. DEPLOYMENT & DEPLOYED SETUP GUIDE ---
    add_h1("7. Production Deployment Architecture")
    doc.add_paragraph(
        "VoxLoop is deployed on an Oracle Cloud Infrastructure (OCI) Always-Free ARM Ampere instance (Ubuntu 22.04 LTS) "
        "containerized using Docker Compose, reverse-proxied via Nginx, and secured with Let's Encrypt SSL (DuckDNS)."
    )

    output_file = Path("c:/Users/USER/Documents/Codex/2026-08-21/verify-upsk-removal/VoxLoop_Detailed_Documentation.docx")
    doc.save(output_file)
    print(f"Documentation successfully created at: {output_file}")

if __name__ == "__main__":
    create_documentation()
